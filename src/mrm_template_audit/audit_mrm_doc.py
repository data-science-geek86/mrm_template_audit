import pandas as pd
import re
import docx

def extract_toc_from_document(file_path):
    """
    Extracts Table of Contents (TOC) entries from a Word document based on paragraph styles.
    
    This function parses the internal XML structure of a .docx file to identify 
    paragraphs styled with 'TOC' variants. It strips out tab characters and page 
    numbers commonly found in Word's generated TOC, then uses regular expressions 
    to separate the heading number from the heading text.

    Args:
        file_path (str): The system path to the .docx file to be processed.

    Returns:
        list[dict]: A list of dictionaries, where each dictionary represents a TOC entry:
            - type (str): Literal description of the entry type ("Literal TOC Entry").
            - level (str): The hierarchical level extracted from the style name (e.g., "1").
            - heading_number (str): The numeric prefix (e.g., "1.1"). Returns "N/A" if not found.
            - heading_text (str): The descriptive title of the section.
            
    Example:
        >>> extract_toc_from_document("report.docx")
        [{'type': 'Literal TOC Entry', 'level': '1', 'heading_number': '1.0', 'heading_text': 'Introduction'}]
    """

    # Load the document
    doc = docx.Document(file_path)
    body_elements = doc._body._body

    # 1. Extract Paragraphs with styles
    ps = body_elements.xpath('.//w:p')
    not_none = [p for p in ps if p.style]
    
    # 2. Filter for TOC styles and remove empty strings
    # We use a list of dictionaries to store the separate components
    toc_data = []

    # Regex pattern to separate Numbering (e.g., 1.2.3) from Title
    # Group 1: The Numbering | Group 2: The Title
    #heading_split_pattern = r'^([\d\.]+)\s+(.*)'
    heading_split_pattern = r'^(?:Section\s+)?(\d+(?:\.\d+)*)\.?\s+(.*)'

    for p in not_none:
        style_name = p.style.lower()
        if style_name.startswith('toc'):
            # Get the level from the style name (e.g., 'toc 1' -> 1)
            level_match = re.search(r'\d+', style_name)
            level = level_match.group() if level_match else "Unknown"
            
            
            # 1. Strip the Page Number (Everything after the LAST tab or large space)
            # We split by tab and keep everything EXCEPT the last element (the page number)
            parts = p.text.split('\t')
            if len(parts) > 1:
                # Rejoin everything except the last part in case titles have tabs
                raw_text = " ".join(parts[:-1]).strip()
            else:
                raw_text = p.text.strip()

            # 2. Robust Regex to separate Numbering from Heading Text
            # Pattern Explanation:
            # ^([\d\.]+)  -> Capture leading digits and dots (Group 1)
            # \.?         -> Optional trailing dot after the number (e.g., '3.')
            # [\s\t]+     -> One or more spaces OR tabs (The separator)
            # (.*)        -> The rest of the text (Group 2)
            heading_split_pattern = r'^([\d\.]+)\.?[\s\t]+(.*)'
            
            if raw_text:
                match = re.match(heading_split_pattern, raw_text)
                if match:
                    # Clean the number (remove any trailing dots like '3.' -> '3')
                    h_number = match.group(1).rstrip('.')
                    h_text = match.group(2).strip()
                else:
                    # Fallback for non-numbered entries (e.g., "Executive Summary")
                    h_number = "N/A"
                    h_text = raw_text

                toc_data.append({
                    "type": "Literal TOC Entry",
                    "level": level,
                    "heading_number": h_number,
                    "heading_text": h_text
                })
    
    return(toc_data)
    
    
    


from docx import Document
import re

def infer_toc_from_document(file_path):
    """
    Infers the Table of Contents (TOC) by analyzing the document's heading hierarchy.

    This function iterates through all paragraphs in a Word document to identify those 
    assigned a 'Heading' style. It extracts the structural level from the style name 
    and uses regular expressions to parse the section numbering and title text from 
    the paragraph content.

    Logic Details:
        - **Style Detection:** Filters paragraphs where the style name starts with 'Heading'.
        - **Level Extraction:** Extracts the digit from the style name (e.g., 'Heading 2' -> level 2).
        - **Numbering Regex:** Specifically handles standard numeric numbering (1.1, 1.2.3) 
          and 'Section' prefixes (e.g., 'Section 1.1 Introduction').
        - **Fallback:** If no numeric prefix is detected, the `heading_number` is set to 'N/A'.

    Args:
        file_path (str): The full path to the .docx file to be analyzed.

    Returns:
        list[dict]: A list of dictionaries representing the inferred structure. 
            Each dictionary contains:
            - **type** (str): The source of the inference ('Document Structure').
            - **level** (str): The hierarchical level of the heading.
            - **heading_number** (str): The extracted section number or 'N/A'.
            - **heading_text** (str): The descriptive title of the heading.

    Example:
        >>> entries = infer_toc_from_document("technical_spec.docx")
        >>> print(entries[0])
        {'type': 'Document Structure', 'level': '1', 'heading_number': '1', 'heading_text': 'Introduction'}
    """
    doc = Document(file_path)
    toc_entries = []
    
    print(f"--- Inferring TOC from Document Structure from: {file_path} ---")
    for para in doc.paragraphs:        
        # 2. Detect Heading Styles (The 'Structure' of the TOC)
        if para.style.name.startswith('Heading'):

            style_name = para.style.name.lower()
            print(style_name)
            #level = para.style.name.split()[-1]
            level_match = re.search(r'\d+', style_name)
            level = level_match.group() if level_match else "Unknown"
            
            # Clean the text: Remove tab characters and page numbers
            # Word TOCs usually look like: "1.1 Section Title\t15"
            raw_text = para.text.split('\t')[0].strip()

            # Regex pattern to separate Numbering (e.g., 1.2.3) from Title
            # Group 1: The Numbering | Group 2: The Title
            heading_split_pattern = r'^(?:Section\s+)?(\d+(?:\.\d+)*)\.?\s+(.*)'
            
            # Separate Numbering from Text
            match = re.match(heading_split_pattern, raw_text)
            print(match)
            if match:
                h_number = match.group(1)
                h_text = match.group(2)
            else:
                # Fallback if no numbering is found (e.g., Executive Summary)
                h_number = "N/A"
                h_text = raw_text
        
            toc_entries.append({
                "type": "Document Structure",
                "level": level,
                "heading_number": h_number,
                "heading_text": h_text
                })
            
    if toc_entries:
        print(f"--- Success - Inferred TOC from Document Structure from: {file_path} ---")
    else:
        print(f"--- Warning - Check the file manually. Unable to infer the TOC from Document Structure from: {file_path} ---")

    return toc_entries


import pandas as pd
from docx import Document
from docx.oxml.ns import qn

def get_numbering_definitions(doc):
    """
    Parses the internal `numbering.xml` of a Word document to extract list start values.

    This function navigates the Word document's XML hierarchy to map specific list 
    instances (`numId`) to their underlying abstract definitions. It identifies the 
    starting integer for every indentation level (ilvl), allowing the parser to 
    respect custom numbering (e.g., a document starting at Section 3 instead of 1).

    Args:
        doc (docx.document.Document): An initialized python-docx Document object.

    Returns:
        dict: A nested dictionary mapping list IDs to level start values.
            Format: `{ numId: { ilvl: start_value } }`
            Example: `{ '1': { 0: 3, 1: 1 } }` indicates a list where the top 
            level starts at 3.

    Note:
        Word stores numbering in two parts: `abstractNum` (the template) and 
        `num` (the instance). This function resolves the link between both.
    """
    try:
        numbering_part = doc.part.numbering_part
        numbering_xml = numbering_part.element
    except (AttributeError, KeyError):
        return {}

    # Define the namespace map for Word
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Map abstractNumId to its levels' start values
    abstract_map = {}
    for abstract_num in numbering_xml.xpath('//w:abstractNum'):
        abs_id = abstract_num.get(qn('w:abstractNumId'))
        levels = {}
        for lvl in abstract_num.xpath('.//w:lvl', namespaces=ns):
            ilvl_val = lvl.get(qn('w:ilvl'))
            if ilvl_val is not None:
                ilvl = int(ilvl_val)
                start_elem = lvl.find(qn('w:start'))
                start_val = int(start_elem.get(qn('w:val'))) if start_elem is not None else 1
                levels[ilvl] = start_val
        abstract_map[abs_id] = levels

    # Map actual numId to the abstract levels
    num_id_map = {}
    for num in numbering_xml.xpath('//w:num'):
        num_id = num.get(qn('w:numId'))
        abs_ref = num.find(qn('w:abstractNumId'))
        if abs_ref is not None:
            abs_id = abs_ref.get(qn('w:val'))
            if abs_id in abstract_map:
                num_id_map[num_id] = abstract_map[abs_id]
                
    return num_id_map

def get_numbering_info(para):
    """
    Extracts numbering metadata (level and list ID) for a specific paragraph.

    The function checks for numbering properties (`numPr`) in two locations:
    1. Directly applied to the paragraph (Direct Formatting).
    2. Defined within the paragraph's associated Style (Style Formatting).

    Args:
        para (docx.text.paragraph.Paragraph): The paragraph object to inspect.

    Returns:
        tuple: A tuple containing (ilvl, numId):
            - ilvl (int): The indentation level (0 for top-level, 1 for sub-level, etc.).
            - numId (str): The unique identifier for the list instance.
            Returns (None, None) if the paragraph is not part of a list.

    Example:
        >>> level, list_id = get_numbering_info(paragraph)
        >>> print(level)
        0
    """
    # Standard namespaces for paragraph xpath
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    numPr = para._p.xpath('.//w:numPr')
    
    if not numPr:
        style_id = para.style.style_id
        styles_xml = para.part.styles.element.xpath(f'./w:style[@w:styleId="{style_id}"]')
        if styles_xml:
            numPr = styles_xml[0].xpath('.//w:numPr')

    if numPr:
        ilvl_elem = numPr[0].find(qn('w:ilvl'))
        numId_elem = numPr[0].find(qn('w:numId'))
        
        ilvl = int(ilvl_elem.get(qn('w:val'))) if ilvl_elem is not None else 0
        numId = numId_elem.get(qn('w:val')) if numId_elem is not None else None
        return ilvl, numId
    
    return None, None



import pandas as pd
from docx import Document

def extract_heading_structure_with_numbering(file_path, num_definitions):
    """
    Extracts the document heading hierarchy by resolving auto-generated Word numbering.

    This function iterates through the paragraphs of a document and identifies those 
    with 'Heading' styles. It uses a stateful counter mechanism to track section 
    numbers across different list instances (numId) and indentation levels (ilvl). 
    It respects custom starting values (e.g., Section 3) by referencing the 
    provided numbering definitions.

    Logic Flow:
        1. **Detection:** Identifies paragraphs where the style name starts with 'Heading'.
        2. **State Tracking:** Uses `active_counters` with a key of `(numId, ilvl)` to 
           support multiple independent lists within a single document.
        3. **Initialization:** If a level is encountered for the first time, it starts 
           at the value defined in `num_definitions`.
        4. **Resetting:** When a higher-level heading is found (e.g., Level 1), all 
           active sub-level counters (Level 2+) for that list are reset to their 
           defined start values.
        5. **Formatting:** Constructs a dot-separated string (e.g., "1.2.1") based 
           on the current state of parent counters.

    Args:
        file_path (str): Path to the .docx file to be processed.
        num_definitions (dict): Numbering metadata obtained from `get_numbering_definitions`.
            Expected format: `{ numId: { ilvl: start_value } }`.

    Returns:
        list[dict]: A list of dictionaries consistent with other extraction tools.
            Each dictionary contains:
            - **heading_number** (str): The resolved section number (e.g., "3.1.2") or "N/A".
            - **heading_text** (str): The text content of the heading.
            - **style** (str): The original Word style name (e.g., "Heading 1").
            - **ilvl** (int or None): The raw indentation level extracted from the XML.

    Example:
        >>> doc = Document("MDD.docx")
        >>> defs = get_numbering_definitions(doc)
        >>> structure = extract_heading_structure_with_numbering("MDD.docx", defs)
    """
    doc = Document(file_path)
    toc_entries = []
    active_counters = {}

    for para in doc.paragraphs:
        # Only process Headings to build TOC
        if para.style.name.startswith('Heading'):
            # Note: Assumes get_numbering_info is defined globally as per previous steps
            ilvl, numId = get_numbering_info(para)
            
            if ilvl is not None and numId is not None:
                state_key = (numId, ilvl)
                
                # 1. Start Value Lookup
                start_val = num_definitions.get(numId, {}).get(ilvl, 1)
                
                # 2. Increment or Initialize
                if state_key not in active_counters:
                    active_counters[state_key] = start_val
                else:
                    active_counters[state_key] += 1
                
                # 3. Reset Children
                for reset_lvl in range(ilvl + 1, 10):
                    child_key = (numId, reset_lvl)
                    if child_key in active_counters:
                        # Reset to start value - 1 so the next increment brings it to start_val
                        active_counters[child_key] = num_definitions.get(numId, {}).get(reset_lvl, 1) - 1

                # 4. Build Number String
                num_parts = []
                for i in range(ilvl + 1):
                    num_parts.append(str(active_counters.get((numId, i), 1)))
                heading_num = ".".join(num_parts)
            else:
                heading_num = "N/A"

            toc_entries.append({
                'heading_number': heading_num,
                'heading_text': para.text.strip(),
                'style': para.style.name,
                'ilvl': ilvl
            })

    return toc_entries



import numpy as np
from docx import Document

def audit_model_document(model_path, template_path):
    """
    Performs a comprehensive audit of a model document against a standard template.
    
    Includes bidirectional checks for missing mandatory sections (template vs model)
    and extra non-standard sections (model vs template).
    """
    audit_report = {
        "internal_discrepancies": [],
        "template_discrepancies": []
    }

    # --- PART 1: Internal TOC vs Body Consistency ---
    literal_toc = extract_toc_from_document(model_path)
    filtered_literal_toc = [
        item for item in literal_toc 
        if not (
            item['type'] == 'Literal TOC Entry' and 
            item['level'] == 'Unknown' and 
            item['heading_number'] == 'N/A' and 
            'contents' in item['heading_text'].lower()
        )
    ]

    model_doc_obj = Document(model_path)
    model_num_defs = get_numbering_definitions(model_doc_obj)
    body_headings = extract_heading_structure_with_numbering(model_path, model_num_defs)

    for i, toc_entry in enumerate(filtered_literal_toc):
        try:
            body_match = body_headings[i]
            if (toc_entry['heading_number'] != body_match['heading_number'] or 
                toc_entry['heading_text'] != body_match['heading_text']):
                
                audit_report["internal_discrepancies"].append({
                    "issue": "TOC and Body Heading Mismatch",
                    "toc_location": f"Entry {i+1}",
                    "details": f"TOC says '{toc_entry['heading_number']} {toc_entry['heading_text']}' "
                               f"but Body contains '{body_match['heading_number']} {body_match['heading_text']}'"
                })
        except IndexError:
            audit_report["internal_discrepancies"].append({
                "issue": "Missing Body Heading",
                "details": f"TOC entry '{toc_entry['heading_text']}' has no corresponding heading in the body."
            })

    # --- PART 2: Model vs Template Compliance ---
    template_structure = extract_toc_from_document(template_path)
    filtered_template_structure = [
        item for item in template_structure 
        if not (
            item['type'] == 'Literal TOC Entry' and 
            item['level'] == 'Unknown' and 
            item['heading_number'] == 'N/A' and 
            'contents' in item['heading_text'].lower()
        )
    ]

    # Create lookup maps (Text -> Number)
    template_map = {item['heading_text'].lower(): item['heading_number'] for item in filtered_template_structure}
    model_map = {item['heading_text'].lower(): item['heading_number'] for item in body_headings}

    # CHECK A: Template vs Model (Missing Mandatory or Incorrect Numbering)
    for t_text, t_num in template_map.items():
        if t_text not in model_map:
            audit_report["template_discrepancies"].append({
                "issue": "Missing Mandatory Section",
                "section": t_text.upper(),
                "details": f"Section '{t_text}' is required by template but missing in model document."
            })
        elif model_map[t_text] != t_num:
            audit_report["template_discrepancies"].append({
                "issue": "Incorrect Section Numbering",
                "section": t_text,
                "details": f"Expected numbering '{t_num}' per template, but found '{model_map[t_text]}' in model."
            })

    # CHECK B: Model vs Template (Extra/Non-Standard Sections)
    for m_text, m_num in model_map.items():
        if m_text not in template_map:
            audit_report["template_discrepancies"].append({
                "issue": "Extra/Non-Standard Section Detected",
                "section": m_text.upper(),
                "details": f"Section '{m_text}' (Number: {m_num}) is not defined in the standard template."
            })

    # Data Attachments for transparency
    audit_report["doc_literal_toc"] = filtered_literal_toc
    audit_report["template_structure"] = filtered_template_structure
    audit_report["body_headings"] = body_headings

    return audit_report


