import pandas as pd
from docx import Document
from docx.oxml.ns import qn

def audit_document_links(doc_path):
    """
    Performs a comprehensive cross-reference audit of hyperlinks and bookmarks in a Word document (.docx).

    Extracts all internal and external hyperlinks as well as bookmarks (including those
    nested inside tables). Cross-references anchors against bookmarks to detect broken links,
    orphaned targets, and resolves the target/source text for internal references.

    Args:
        doc_path (str): File path to the Word document (.docx) to be audited.

    Returns:
        pd.DataFrame: A Pandas DataFrame containing the audit report with columns:
            - 'Link Text': Display text of the hyperlink or placeholder for unlinked targets.
            - 'Type': Category of the reference ('Internal (Bookmark)', 'External (URL)',
              or 'Orphaned Bookmark').
            - 'Target ID/URL': The bookmark name or external target URL.
            - 'Target Text': The text inside the targeted bookmark range, or contextual text
              if pointing to a zero-width point bookmark.
            - 'Audit Status': Validation state ('Working', 'BROKEN (Missing Target)',
              'Active (External)', or 'Warning (Unused Target)').

    Notes:
        - Automatically excludes Word built-in hidden bookmarks (e.g., '_Toc', '_Ref', '_GoBack').
        - Traverses both standard document body paragraphs and table cells.
        - Verification of external URL reachability requires network HTTP requests 
          and is currently tagged as 'Active (External)' based solely on relation presence.
    """
    document = Document(doc_path)
    doc_element = document.part._element
    rels = document.part.rels
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # --- Helper: Extract text for a given bookmarkStart node ---
    def get_bookmark_text(start_node):
        bookmark_id = start_node.get(qn('w:id'))
        source_text = ""
        
        # 1. Attempt to get range text (text between start and end tags)
        curr = start_node.getnext()
        while curr is not None:
            if curr.tag == qn('w:bookmarkEnd') and curr.get(qn('w:id')) == bookmark_id:
                break
            if curr.tag == qn('w:r'):
                for t in curr.xpath('.//w:t', namespaces=ns):
                    if t.text:
                        source_text += t.text
            curr = curr.getnext()

        # 2. Fallback for Point Bookmarks: extract parent paragraph text
        if not source_text.strip():
            parent_paragraph = start_node.getparent()
            if parent_paragraph is not None:
                full_par_text = "".join([t.text for t in parent_paragraph.xpath('.//w:t') if t.text])
                source_text = f"[Point at: {full_par_text[:50].strip()}...]" if full_par_text.strip() else "[Empty Location]"
            else:
                source_text = "[Empty Location]"

        return source_text.strip()

    # --- 1. Extract Valid Bookmark Targets & Text ---
    # Map structure: { bookmark_name: {"utilized": False, "target_text": "..."} }
    valid_bookmarks = {}
    for start_node in doc_element.xpath('.//w:bookmarkStart'):
        name = start_node.get(qn('w:name'))
        if name and not name.startswith(('_Toc', '_Ref', '_GoBack')):
            valid_bookmarks[name] = {
                "utilized": False,
                "target_text": get_bookmark_text(start_node)
            }

    # --- Helper: Gather all paragraphs (Body + Tables) ---
    def iter_all_paragraphs(doc):
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    # --- 2. Extract and Validate Hyperlinks ---
    link_audit_log = []
    
    for p in iter_all_paragraphs(document):
        hyperlinks = p._element.findall('.//' + qn('w:hyperlink'))
        for href in hyperlinks:
            display_text = "".join([t.text for t in href.findall('.//' + qn('w:t')) if t.text])
            r_id = href.get(qn('r:id'))
            anchor = href.get(qn('w:anchor'))
            
            status = "Unknown"
            link_type = "Internal"
            target = "N/A"
            target_text = "N/A"

            if r_id and r_id in rels:
                target = rels[r_id].target_ref
                link_type = "External (URL)"
                status = "Unknown"
                target_text = "[External Web Location]"
            elif anchor:
                target = anchor
                link_type = "Internal (Bookmark)"
                if anchor in valid_bookmarks:
                    status = "Working"
                    valid_bookmarks[anchor]["utilized"] = True  # Mark bookmark as used
                    target_text = valid_bookmarks[anchor]["target_text"]
                else:
                    status = "BROKEN (Missing Target)"
                    target_text = "[Target Not Found]"
            
            link_audit_log.append({
                "Link Text": display_text.strip(),
                "Type": link_type,
                "Target ID/URL": target,
                "Target Text": target_text,
                "Audit Status": status
            })

    # --- 3. Identify Orphaned Bookmarks ---
    for b_name, b_info in valid_bookmarks.items():
        if not b_info["utilized"]:
            link_audit_log.append({
                "Link Text": "[N/A - Unlinked Target]",
                "Type": "Orphaned Bookmark",
                "Target ID/URL": b_name,
                "Target Text": b_info["target_text"],
                "Audit Status": "Warning (Unused Target)"
            })

    return pd.DataFrame(link_audit_log)
